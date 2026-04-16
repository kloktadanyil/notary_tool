import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from placeholders import *

def replace_placeholder_in_document(document_obj, data_to_fill):
    """
    Replaces all placeholders in the document, preserving the formatting of the surrounding text.
    """
    # for paragraph
    for paragraph in document_obj.paragraphs:
        for ph, value in data_to_fill.items():
            if ph in paragraph.text:
                for run in paragraph.runs:
                    if ph in run.text:
                        # Changing text
                        run.text = run.text.replace(ph, str(value))
                        
    # For table
    for table in document_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for ph, value in data_to_fill.items():
                        if ph in p.text:
                            # changing text
                            p.text = p.text.replace(ph, str(value))
                    
                    
def remove_text_between_placeholders(document_obj, start_placeholder, end_placeholder, replacement=""):
    """
    Deletes part of the text in a paragraph from one placeholder to another.
    """
    for paragraph in document_obj.paragraphs:
        if start_placeholder in paragraph.text and end_placeholder in paragraph.text:
            start_index = paragraph.text.find(start_placeholder)
            end_index = paragraph.text.find(end_placeholder) + len(end_placeholder)
            
            text_before = paragraph.text[:start_index]
            text_after = paragraph.text[end_index:]
            
            paragraph.text = text_before + replacement + text_after

def convert_digit_to_text_uk(n, case="ordinal"):
    """
    Converts an integer to Ukrainian text format, supporting
    quantitative, ordinal, and genitive cases.

    :param n: Number to convert (expected up to 9999).
    :param case: 'regular', 'ordinal' (ordinal - which?), 'genitive' (genitive - who/what?).
    :return: Text representation of the number.
    """
    # for the Ordinal Case 
    ones_ordinal = ["", "перше", "друге", "третє", "четверте", "п’яте", "шосте", "сьоме", "восьме", "дев’яте"]
    teens_ordinal = ["", "одинадцяте", "дванадцяте", "тринадцяте", "чотирнадцяте", "п’ятнадцяте",
                     "шістнадцяте", "сімнадцяте", "вісімнадцяте", "дев’ятнадцяте"]
    tens_ordinal = ["", "десяте", "двадцяте", "тридцяте", "сорокове", "п’ятдесяте",
                    "шістдесяте", "сімдесяте", "вісімдесяте", "дев’яносте"]
    
    # for Genitive ---
    ones_genitive = ["", "першого", "другого", "третього", "четвертого", "п’ятого", "шостого", "сьомого", "восьмого", "дев’ятого"]
    teens_genitive = ["", "одинадцятого", "дванадцятого", "тринадцятого", "чотирнадцятого", "п’ятнадцятого",
                      "шістнадцятого", "сімнадцятого", "вісімнадцятого", "дев’ятнадцятого"]
    tens_genitive = ["", "десятого", "двадцятого", "тридцятого", "сорокового", "п’ятдесятого",
                      "шістдесятого", "сімдесятого", "вісімдесятого", "дев’яностого"]

    # -for Regular case ---
    ones_regular = ["", "один", "два", "три", "чотири", "п’ять", "шість", "сім", "вісім", "дев’ять"]
    teens_regular = ["", "одинадцять", "дванадцять", "тринадцять", "чотирнадцять", "п’ятнадцять",
                     "шістнадцять", "сімнадцять", "вісімнадцять", "дев’ятнадцять"]
    tens_regular = ["", "десять", "двадцять", "тридцять", "сорок", "п’ятдесят",
                    "шістдесят", "сімдесят", "вісімдесят", "дев’яносто"]
    
    hundreds = ["", "сто", "двісті", "триста", "чотириста", "п’ятсот",
                "шістсот", "сімсот", "вісімсот", "дев’ятсот"]
    
    result = []
    
    if n == 0:
        return ""
        
    original_n = n
    
    # 1. Changing thousands
    if n >= 1000:
        thousand_part = n // 1000
        n %= 1000
        thousands_for_declension = thousand_part 
        
        # Changing 100-999 
        if thousand_part >= 100:
            result.append(hundreds[thousand_part // 100])
            thousand_part %= 100
            
        if 11 <= thousand_part <= 19:
            result.append(teens_regular[thousand_part - 10])
        elif thousand_part >= 20:
            result.append(tens_regular[thousand_part // 10])
            if thousand_part % 10 > 0:
                # Changing 1-2
                if thousand_part % 10 == 1:
                    result.append("одна") # 1001, 2001
                elif thousand_part % 10 == 2:
                    result.append("дві")
                else:
                    result.append(ones_regular[thousand_part % 10])
        elif thousand_part > 0:
            if thousand_part == 1:
                result.append("одна")
            elif thousand_part == 2:
                result.append("дві")
            else:
                result.append(ones_regular[thousand_part])
            
        # Ending for "тисяча"
        last_two_digits_of_thousands = thousands_for_declension % 100
        last_digit_of_thousands = thousands_for_declension % 10

        if 11 <= last_two_digits_of_thousands <= 14:
             result.append("тисяч")
        elif last_digit_of_thousands == 1:
            result.append("тисяча")
        elif 2 <= last_digit_of_thousands <= 4:
            result.append("тисячі")
        else:
            result.append("тисяч")
            
    # 2. changing 0-999
    if n > 0:
        # Hundreds quantitative
        hundreds_part = n // 100
        n %= 100
        
        if hundreds_part > 0:
            result.append(hundreds[hundreds_part])
            
        # 0-99
        
        # for Ordinal Case
        if case == "ordinal":
            if n == 0:
                pass
            elif n % 10 == 0 and n >= 20:
                result.append(tens_ordinal[n // 10])
            elif n >= 20:
                result.append(tens_regular[n // 10]) 
                n %= 10
                if n > 0:
                    result.append(ones_ordinal[n])
            elif n >= 11:
                result.append(teens_ordinal[n - 10])
            elif n > 0:
                result.append(ones_ordinal[n])
        
        # for Genitive Case
        elif case == "genitive":
            if n == 0:
                pass 
            elif n % 10 == 0 and n >= 20:
                # 20, 30: 'двадцятого'
                result.append(tens_genitive[n // 10])
            elif n >= 20:
                # 25: 'двадцять' + 'п’ятого'
                result.append(tens_regular[n // 10]) 
                n %= 10
                if n > 0:
                    result.append(ones_genitive[n])
            elif n >= 11:
                # 11-19: 'одинадцятого'
                result.append(teens_genitive[n - 10])
            elif n > 0:
                #  1-10: 'першого', 'десятого'
                result.append(ones_genitive[n])
                
        # for Regular Case
        elif case == "regular":
             if 11 <= n <= 19:
                 result.append(teens_regular[n - 10])
             elif n >= 20:
                 result.append(tens_regular[n // 10])
                 n %= 10
                 if n > 0:
                     result.append(ones_regular[n])
             elif n > 0:
                 result.append(ones_regular[n])

    return " ".join(result).strip()

def get_current_date_uk_full():
    from datetime import date
    today = date.today()
    
    # 1. Day for Ordinal case (яке?)
    day_text = convert_digit_to_text_uk(today.day, case="ordinal")
    
    months_uk = [
        "січня", "лютого", "березня", "квітня", "травня", "червня",
        "липня", "серпня", "вересня", "жовтня", "листопада", "грудня"
    ]
    # 2. Month: Genitive
    month_text = months_uk[today.month - 1]
    
    # 3. Year: Genitive (кого? чого? - року)
    year_text = convert_digit_to_text_uk(today.year, case="genitive")
    
    return f"{day_text} {month_text} {year_text}"

def get_current_date_uk_full():
    from datetime import date
    today = date.today()
    
    # 1. День: Порядковий відмінок (яке?)
    day_text = convert_digit_to_text_uk(today.day, case="ordinal")
    
    months_uk = [
        "січня", "лютого", "березня", "квітня", "травня", "червня",
        "липня", "серпня", "вересня", "жовтня", "листопада", "грудня"
    ]
    # 2. Місяць: Родовий відмінок
    month_text = months_uk[today.month - 1]
    
    # 3. Рік: Родовий відмінок (кого/чого? - року)
    year_text = convert_digit_to_text_uk(today.year, case="genitive")
    
    return f"{day_text} {month_text} {year_text}"
    
def generate_children_text(children_list):
    """
    Генерує єдиний рядок з даними про всіх дітей для вставки в документ.
    """
    if not children_list:
        return ""
    
    children_strings = []
    
    for child in children_list:
        pib = child.get('pib', '')
        dob = child.get('dob', '')
        age_status = child.get('age_status', '')
        gender_ending = child.get('gender_ending', '')
        ending_1 = child.get('ending_1', '')
        ending_2 = child.get('ending_2', '')

        # Форматуємо рядок для однієї дитини
        text = f"мо{ending_1} {age_status}{ending_2} {gender_ending} {pib}, {dob} року народження,"
        children_strings.append(text)

    # Об'єднуємо всі рядки. Якщо дітей більше однієї, додаємо "та"
    if len(children_strings) == 1:
        return children_strings[0]
    else:
        return " та/або ".join(children_strings)

def generate_children_text_2 (children_list):
    
    if not children_list:
        return "",""
    
    if len(children_list) == 1:
            zaymennyk_ditey = "моєї дитини,"
            my_child =  "мою дитину"
            my_child_2 =  'життя моєї дитини на період її'
            
    else:
            zaymennyk_ditey = "моїх дітей,"
            my_child =  "моїх дітей"
            my_child_2 =  'життя моїх дітей на період їх'
    children_strings_in_the_end = []
    
    for child in children_list:
        pib = child.get('pib', '')
    
        text = f" {pib} "
        children_strings_in_the_end.append(text)
    if len(children_strings_in_the_end) == 1:
            return zaymennyk_ditey + children_strings_in_the_end[0], my_child, my_child_2
    else:
        return zaymennyk_ditey + "та".join(children_strings_in_the_end), my_child, my_child_2
        
    
def generate_companions_text(companion_list):
    """
    Генерує єдиний рядок з даними про всіх супроводжуючих для вставки в документ.
    Використовує кому після дати народження та сполучник 'та/або' для об'єднання.
    """
    if not companion_list:
        return "", "", ""
    
    companion_strings = []
    
    for companion in companion_list:
        pib = companion.get('pib', '')
        dob = companion.get('dob', '')
        gender_ending_suprov = companion.get('gender_ending_suprov', '')

        # Форматуємо рядок для одного супроводжуючого
        # Додаємо кому в кінці рядка, як було вказано.
        text = f"{pib}, {dob} року народження,"
        companion_strings.append(text)


    num_companions = len(companion_strings)
    if num_companions == 1:
        companion_ending = f"{gender_ending_suprov} бере"
        companion_ending_2 = "зобов\'язується"
    else:
        companion_ending = "які беруть"
        companion_ending_2 = "зобов\'язуються"
    
    if num_companions == 1:
        # У цьому випадку кома вже є в кінці, але ми її прибрали. Потрібно вирішити,
        # чи потрібна кома в документі після єдиного супроводжуючого. 
        # Залишимо без коми, генератор додасть її, якщо потрібно.
        return companion_strings[0], companion_ending, companion_ending_2
    
    elif num_companions >= 2:
        return " та/або ".join(companion_strings), companion_ending, companion_ending_2
    
    return ""

# --- Основна функція для генерації документа ---
def generate_document_from_data(data_to_fill, status_label):
    try:
        template_path = 'Templates.docx'
        pib = data_to_fill.get('{{ПІБ}}')
        # Прописуємо шлях куди зберігається документ
        output_dir = os.path.join('Згенеровані документи', 'Заяви','За кордон') 
        filename = f"{pib.replace(' ', '_')}.docx"
        output_path = os.path.join(output_dir, filename)
        # Дозволяє створювати папки
        os.makedirs(output_dir, exist_ok=True)
        document = Document(template_path)
        
        # Генеруємо текст про дітей і додаємо його до словника даних 
        children_list = data_to_fill.get("children", [])
        data_to_fill[CHILDREN_DATA] = generate_children_text(children_list)
        data_to_fill[CHILDREN_DATA_2],data_to_fill[MY_CHILD],data_to_fill[MY_CHILD_2]= generate_children_text_2(children_list)
        print(f"Значення CHILDREN_DATA_2: {data_to_fill.get(CHILDREN_DATA_2)}")
        # Генеруємо текст про супроводжуючих і додаємо його до словника даних
        companions_list = data_to_fill.get("companions", [])
        data_to_fill[COMPANIONS_DATA], data_to_fill[COMPANIONS_ENDING], data_to_fill[COMPANIONS_ENDING_2] = generate_companions_text(companions_list)
        data_to_fill[DATE_PROPYSOM] = get_current_date_uk_full ()
        

        # --- ЛОГІКА ВИДАЛЕННЯ БЛОКІВ ТЕПЕР ТУТ ---
        # Ми перевіряємо, які дані є у словнику data_to_fill
        if '{{ПАСПОРТ}}' in data_to_fill:
            # Якщо є дані для старого паспорта, видаляємо блок нового
            remove_text_between_placeholders(document, '{{БЛОК_НОВИЙ_ПАСПОРТ_СТАРТ}}', '{{БЛОК_НОВИЙ_ПАСПОРТ_КІНЕЦЬ}}')
        elif '{{НОМЕР_КАРТКИ}}' in data_to_fill:
            # Якщо є дані для нового паспорта, видаляємо блок старого
            remove_text_between_placeholders(document, '{{БЛОК_СТАРИЙ_ПАСПОРТ_СТАРТ}}', '{{БЛОК_СТАРИЙ_ПАСПОРТ_КІНЕЦЬ}}')

        # Заміна всіх плейсхолдерів (включно з маркерами блоків, які залишились)
        replace_placeholder_in_document(document, data_to_fill)
            
        document.save(output_path)
        status_label.config(text=f"Документ збережено як '{output_path}'!", style="Green.TLabel")
    
    except FileNotFoundError:
        status_label.config(text=f"ПОМИЛКА: Файл '{template_path}' не знайдено.", style="Red.TLabel")
    except Exception as e:
        status_label.config(text=f"Виникла помилка: {e}", style="Red.TLabel")
